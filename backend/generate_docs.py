import os
import urllib.request
import json
from docx import Document

# Base paths
DOCS_DIR = "/app/docs/feature"
AGENTS_DIR = os.path.join(DOCS_DIR, "agents")
TOOLS_DIR = os.path.join(DOCS_DIR, "tools")

os.makedirs(AGENTS_DIR, exist_ok=True)
os.makedirs(TOOLS_DIR, exist_ok=True)

# Fetch Agents
agents_req = urllib.request.urlopen('http://localhost:8000/api/agents')
agents = json.loads(agents_req.read())

# Fetch Tools
tools_req = urllib.request.urlopen('http://localhost:8000/api/tools')
tools = json.loads(tools_req.read())

def create_agent_md(agent):
    content = f"""# {agent['name']}

## Overview
**Type:** {agent['agent_type']}
**Description:** {agent['description']}

## Configuration
- **Max Steps:** {agent['max_steps']}
- **Timeout:** {agent['timeout_seconds']}s
- **Memory Enabled:** {agent['memory_enabled']}
- **Tools Permitted:** {', '.join(agent['tools_enabled']) if agent['tools_enabled'] else 'None'}

## Technical Architecture
The agent is orchestrated using the `base.py` execution loop. It binds dynamically to the specified tool profiles and enforces maximum iteration caps to prevent infinite loops.

## Demos

### Demo 1: Basic Inquiry
**User:** Can you help me with a basic task?
**Agent:** Yes, I can assist you based on my configured tools.

### Demo 2: Tool Execution
**User:** Use your tool to retrieve data.
**Agent:** Accessing tool... Data retrieved successfully.

### Demo 3: Edge Case
**User:** Do something outside your scope.
**Agent:** I apologize, but I am not authorized to perform that action.
"""
    with open(os.path.join(AGENTS_DIR, f"{agent['slug']}.md"), "w", encoding="utf-8") as f:
        f.write(content)

def create_agent_docx(agent):
    doc = Document()
    doc.add_heading(agent['name'], 0)
    
    doc.add_heading('Overview', level=1)
    doc.add_paragraph(f"Type: {agent['agent_type']}")
    doc.add_paragraph(f"Description: {agent['description']}")
    
    doc.add_heading('Demos', level=1)
    doc.add_paragraph('Demo 1: Basic usage example showing how the agent responds.')
    doc.add_paragraph('Demo 2: Example of the agent using a tool.')
    doc.add_paragraph('Demo 3: Example of the agent gracefully handling an error.')
    
    doc.save(os.path.join(AGENTS_DIR, f"{agent['slug']}.docx"))

def create_tool_md(tool):
    content = f"""# {tool['name']}

## Overview
**Category:** {tool['category']}
**Description:** {tool['description']}

## Schema Details
- **Safe Mock Mode:** {tool['safe_mock_mode']}
- **Requires Auth:** {tool['requires_auth']}
- **Required Env Keys:** {', '.join(tool['required_env_keys']) if tool['required_env_keys'] else 'None'}

### Input Schema
```json
{json.dumps(tool['input_schema'], indent=2)}
```

### Output Schema
```json
{json.dumps(tool['output_schema'], indent=2)}
```
"""
    with open(os.path.join(TOOLS_DIR, f"{tool['slug']}.md"), "w", encoding="utf-8") as f:
        f.write(content)

def create_tool_docx(tool):
    doc = Document()
    doc.add_heading(tool['name'], 0)
    
    doc.add_heading('Overview', level=1)
    doc.add_paragraph(f"Category: {tool['category']}")
    doc.add_paragraph(f"Description: {tool['description']}")
    
    doc.add_heading('Usage', level=1)
    doc.add_paragraph("This tool is designed to be called by an AI agent to perform specific actions securely.")
    
    doc.save(os.path.join(TOOLS_DIR, f"{tool['slug']}.docx"))

# Generate Master Guide
def create_master_guide():
    content = f"""# AgentHive Master Feature Guide

## Agents ({len(agents)})
{chr(10).join([f"- {a['name']} ({a['slug']})" for a in agents])}

## Tools ({len(tools)})
{chr(10).join([f"- {t['name']} ({t['slug']})" for t in tools])}

## Workflow Examples
- Multi-agent orchestration
- Cross-platform data syncing
"""
    with open(os.path.join(DOCS_DIR, "MASTER_GUIDE.md"), "w", encoding="utf-8") as f:
        f.write(content)
        
    doc = Document()
    doc.add_heading("AgentHive Master Feature Guide", 0)
    doc.add_paragraph(f"Total Agents: {len(agents)}")
    doc.add_paragraph(f"Total Tools: {len(tools)}")
    doc.save(os.path.join(DOCS_DIR, "MASTER_GUIDE.docx"))

# Execute
print("Generating Agents documentation...")
for agent in agents:
    create_agent_md(agent)
    create_agent_docx(agent)

print("Generating Tools documentation...")
for tool in tools:
    create_tool_md(tool)
    create_tool_docx(tool)

print("Generating Master Guide...")
create_master_guide()

print("Documentation generated successfully.")
